"""DOCX (Word) extractor — returns one span per non-empty paragraph."""

from __future__ import annotations

import io

from app.ingest_docs.extractors.pdf_text import ExtractedSpan


def extract_docx(body: bytes) -> list[ExtractedSpan]:
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError as exc:  # pragma: no cover — install [parse] extras
        raise RuntimeError("python-docx not installed (pip install 'cpa[parse]')") from exc

    doc = Document(io.BytesIO(body))
    spans: list[ExtractedSpan] = []
    seq = 0
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        seq += 1
        spans.append(ExtractedSpan(anchor=f"paragraph={seq}", text=text))
    # Tables can carry policy text too. Walk every cell and append.
    for t_idx, table in enumerate(doc.tables, start=1):
        for r_idx, row in enumerate(table.rows, start=1):
            row_text = " | ".join((cell.text or "").strip() for cell in row.cells if (cell.text or "").strip())
            if row_text:
                spans.append(ExtractedSpan(anchor=f"table={t_idx}.row={r_idx}", text=row_text))
    return spans
